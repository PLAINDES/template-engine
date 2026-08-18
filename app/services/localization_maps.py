# app/services/localization_maps.py
from __future__ import annotations

import hashlib
import json
import math
from io import BytesIO
from pathlib import Path

import httpx
from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from minio import Minio
from minio.error import S3Error

from app.core.config import settings
from app.utils.minio_client import client as minio_prosedi, upload_to_minio

CACHE_DIR = Path("/tmp/prosedi-mapas")
CACHE_DIR.mkdir(parents=True, exist_ok=True)

COLOR_FONDO = (248, 250, 252)
COLOR_UNIDAD = (216, 226, 234)
COLOR_BORDE = (120, 144, 164)
COLOR_RESALTE = (244, 195, 139)
COLOR_RESALTE_BORDE = (194, 105, 31)
COLOR_PIN = (220, 38, 38)

_geo_cache: dict[str, list] = {}


def insertar_mapas_de_localizacion(docx_bytes: bytes, values: dict[str, str]) -> bytes:
    departamento = _texto(values, "NOMBRE_DEPARTAMENTO")
    provincia = _texto(values, "NOMBRE_PROVINCIA")
    distrito = _texto(values, "NOMBRE_DISTRITO")
    lat = _float(values, "LATITUD_IE")
    lon = _float(values, "LONGITUD_IE")

    if not departamento or not provincia or not distrito:
        return docx_bytes

    mapas = generar_mapas(
        departamento=departamento,
        provincia=provincia,
        distrito=distrito,
        lat=lat,
        lon=lon,
    )

    doc = Document(BytesIO(docx_bytes))
    tabla = _tabla_de_localizacion(doc)
    if tabla is None:
        return docx_bytes

    _poner_imagen(tabla.rows[1].cells[0], mapas["peru"])
    _poner_imagen(tabla.rows[1].cells[1], mapas["departamento"])
    _poner_imagen(tabla.rows[3].cells[0], mapas["provincia"])
    _poner_imagen(tabla.rows[3].cells[1], mapas["satelite"])

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


def generar_mapas(
    departamento: str,
    provincia: str,
    distrito: str,
    lat: float | None,
    lon: float | None,
) -> dict[str, bytes]:
    clave = _clave_cache(departamento, provincia, distrito, lat, lon)
    nombres = ("peru", "departamento", "provincia", "satelite")
    rutas = {nombre: CACHE_DIR / f"{clave}-{nombre}.png" for nombre in nombres}

    if all(p.exists() for p in rutas.values()):
        mapas = {k: p.read_bytes() for k, p in rutas.items()}
        _guardar_pngs_minio(clave, mapas)
        return mapas

    desde_minio = {nombre: _leer_png_minio(clave, nombre) for nombre in nombres}
    if all(png is not None for png in desde_minio.values()):
        listos = {nombre: png for nombre, png in desde_minio.items() if png is not None}
        for nombre, png in listos.items():
            rutas[nombre].write_bytes(png)
        return listos

    departamentos = _cargar_cords("departamentosCords.json")
    provincias = _cargar_cords("provinciasCords.json")
    distritos = _cargar_cords("distritosCords.json")

    peru_png = _dibujar_mapa(
        unidades=departamentos,
        nombre_campo="departamento",
        resalte=departamento,
        pin=(lon, lat) if lat is not None and lon is not None else None,
    )
    depto_png = _dibujar_mapa(
        unidades=[
            u
            for u in provincias
            if _igual(u.get("departamento"), departamento)
        ],
        nombre_campo="provincias",
        resalte=provincia,
        pin=(lon, lat) if lat is not None and lon is not None else None,
    )
    prov_png = _dibujar_mapa(
        unidades=[
            u
            for u in distritos
            if _igual(u.get("departamento"), departamento)
            and _igual(u.get("provincias"), provincia)
        ],
        nombre_campo="distrito",
        resalte=distrito,
        pin=(lon, lat) if lat is not None and lon is not None else None,
    )
    sat_png = _satelite(lat, lon)

    mapas = {
        "peru": peru_png,
        "departamento": depto_png,
        "provincia": prov_png,
        "satelite": sat_png,
    }
    for nombre, png in mapas.items():
        rutas[nombre].write_bytes(png)
    _guardar_pngs_minio(clave, mapas)
    return mapas


def _png_minio_key(clave: str, nombre: str) -> str:
    return f"mapas-localizacion/{clave}/{nombre}.png"


def _guardar_pngs_minio(clave: str, mapas: dict[str, bytes]) -> None:
    for nombre, png in mapas.items():
        if _existe_png_minio(clave, nombre):
            continue
        upload_to_minio(
            _png_minio_key(clave, nombre),
            png,
            content_type="image/png",
        )


def _existe_png_minio(clave: str, nombre: str) -> bool:
    try:
        minio_prosedi.stat_object(settings.MINIO_BUCKET, _png_minio_key(clave, nombre))
        return True
    except S3Error as e:
        if e.code in ("NoSuchKey", "NoSuchBucket"):
            return False
        raise


def _leer_png_minio(clave: str, nombre: str) -> bytes | None:
    key = _png_minio_key(clave, nombre)
    try:
        respuesta = minio_prosedi.get_object(settings.MINIO_BUCKET, key)
        try:
            return respuesta.read()
        finally:
            respuesta.close()
            respuesta.release_conn()
    except S3Error as e:
        if e.code in ("NoSuchKey", "NoSuchBucket"):
            return None
        raise


def _satelite(lat: float | None, lon: float | None) -> bytes:
    if lat is None or lon is None:
        raise ValueError(
            "La IE no tiene latitud/longitud; no se puede armar la vista satelital."
        )
    delta = 0.006
    bbox = f"{lon - delta},{lat - delta},{lon + delta},{lat + delta}"
    url = (
        "https://services.arcgisonline.com/ArcGIS/rest/services/"
        "World_Imagery/MapServer/export"
        f"?bbox={bbox}&bboxSR=4326&imageSR=4326&size=800,800&format=png&f=image"
    )
    respuesta = httpx.get(url, timeout=30.0)
    if respuesta.status_code != 200:
        raise ValueError(
            f"La vista satelital no se pudo descargar (HTTP {respuesta.status_code})."
        )
    img = Image.open(BytesIO(respuesta.content)).convert("RGBA")
    draw = ImageDraw.Draw(img)
    cx, cy = img.size[0] // 2, img.size[1] // 2
    r = 10
    draw.ellipse((cx - r, cy - r, cx + r, cy + r), fill=COLOR_PIN)
    draw.ellipse(
        (cx - r - 3, cy - r - 3, cx + r + 3, cy + r + 3),
        outline=(255, 255, 255),
        width=3,
    )
    out = BytesIO()
    img.convert("RGB").save(out, format="PNG")
    return out.getvalue()


def _dibujar_mapa(
    unidades: list,
    nombre_campo: str,
    resalte: str,
    pin: tuple[float, float] | None,
    size: int = 800,
) -> bytes:
    if not unidades:
        raise ValueError(f"No hay geometría para resaltar {resalte}")

    anillos = []
    for u in unidades:
        for anillo in _anillos(u.get("coordenadas")):
            anillos.append((anillo, _igual(u.get(nombre_campo), resalte)))
    if not anillos:
        raise ValueError(f"Sin coordenadas para {resalte}")

    xs = [p[0] for anillo, _ in anillos for p in anillo]
    ys = [p[1] for anillo, _ in anillos for p in anillo]
    minx, maxx = min(xs), max(xs)
    miny, maxy = min(ys), max(ys)
    pad = max(maxx - minx, maxy - miny) * 0.06
    minx, maxx = minx - pad, maxx + pad
    miny, maxy = miny - pad, maxy + pad

    img = Image.new("RGB", (size, size), COLOR_FONDO)
    draw = ImageDraw.Draw(img)

    def xy(lon: float, lat: float) -> tuple[int, int]:
        x = int((lon - minx) / (maxx - minx) * (size - 1))
        y = int((1 - (lat - miny) / (maxy - miny)) * (size - 1))
        return x, y

    for anillo, es_resalte in anillos:
        pts = [xy(lon, lat) for lon, lat in _simplificar(anillo)]
        if len(pts) < 3:
            continue
        fill = COLOR_RESALTE if es_resalte else COLOR_UNIDAD
        borde = COLOR_RESALTE_BORDE if es_resalte else COLOR_BORDE
        draw.polygon(pts, fill=fill, outline=borde)

    if pin is not None:
        px, py = xy(pin[0], pin[1])
        r = 7
        draw.ellipse((px - r, py - r, px + r, py + r), fill=COLOR_PIN)

    out = BytesIO()
    img.save(out, format="PNG")
    return out.getvalue()


def _cargar_cords(nombre: str) -> list:
    if nombre in _geo_cache:
        return _geo_cache[nombre]
    cliente = Minio(
        endpoint=f"{settings.MINIO_ENDPOINT}:{settings.MINIO_PORT}",
        access_key=settings.MINIO_ACCESS_KEY,
        secret_key=settings.MINIO_SECRET_KEY,
        secure=settings.MINIO_USE_SSL,
    )
    respuesta = cliente.get_object("plaindes", f"cords/{nombre}")
    try:
        data = json.loads(respuesta.read().decode("utf-8"))
    finally:
        respuesta.close()
        respuesta.release_conn()
    if not isinstance(data, list):
        raise ValueError(f"{nombre} no es una lista de geometrías")
    _geo_cache[nombre] = data
    return data


def _anillos(coordenadas) -> list[list[tuple[float, float]]]:
    puntos: list[list[tuple[float, float]]] = []

    def caminar(nodo) -> None:
        if not isinstance(nodo, list) or not nodo:
            return
        if isinstance(nodo[0], (int, float)):
            return
        if (
            isinstance(nodo[0], list)
            and len(nodo[0]) >= 2
            and isinstance(nodo[0][0], (int, float))
        ):
            puntos.append([(float(p[0]), float(p[1])) for p in nodo if len(p) >= 2])
            return
        for hijo in nodo:
            caminar(hijo)

    caminar(coordenadas)
    return puntos


def _simplificar(anillo: list[tuple[float, float]], max_puntos: int = 1200):
    if len(anillo) <= max_puntos:
        return anillo
    paso = math.ceil(len(anillo) / max_puntos)
    recorte = anillo[::paso]
    if recorte[-1] != anillo[-1]:
        recorte.append(anillo[-1])
    return recorte


def _tabla_de_localizacion(doc: Document):
    for tabla in doc.tables:
        if len(tabla.rows) < 4 or len(tabla.columns) < 2:
            continue
        texto = tabla.rows[0].cells[0].text.strip().lower()
        if "macro localizaci" in texto:
            return tabla
    return None


def _poner_imagen(celda, png: bytes) -> None:
    primero = celda.paragraphs[0]
    nuevo = OxmlElement("w:p")
    primero._element.addprevious(nuevo)
    parrafo = Paragraph(nuevo, primero._parent)
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = parrafo.add_run()
    run.add_picture(BytesIO(png), width=Cm(7.2))
    for p in celda.paragraphs:
        for r in p.runs:
            r.font.size = Pt(8)


def _texto(values: dict[str, str], key: str) -> str:
    return str(values.get(key) or "").strip()


def _float(values: dict[str, str], key: str) -> float | None:
    crudo = _texto(values, key)
    if crudo == "":
        return None
    return float(crudo.replace(",", "."))


def _igual(a, b: str) -> bool:
    return str(a or "").strip().casefold() == b.strip().casefold()


def _clave_cache(
    departamento: str,
    provincia: str,
    distrito: str,
    lat: float | None,
    lon: float | None,
) -> str:
    base = f"{departamento}|{provincia}|{distrito}|{lat}|{lon}"
    return hashlib.sha1(base.encode("utf-8")).hexdigest()[:16]
