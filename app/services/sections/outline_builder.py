# app/services/sections/outline_builder.py
"""
Índice completo del documento: todos los títulos con sus variables y el
marcador que permite saltar a cada uno desde el editor.

A diferencia de `extract_section_structure`, que devuelve el árbol de una sola
sección H1, aquí se recorre el documento entero de una vez — es lo que pinta el
índice lateral de la plataforma.
"""
from io import BytesIO
from typing import Any, Dict, List

from docx import Document
from docx.table import Table

from app.services.parser import VAR_RE, TABLE_KEYS, IMAGE_KEYS
from app.services.sections.heading_parser import (
    _get_heading_level,
    refine_level_with_numbering,
)
from app.services.variable_marker import heading_bookmark


def _collect_variables(text: str, node: Dict[str, Any]) -> None:
    """Añade al nodo las variables del texto, sin repetir y sin tablas/imágenes.

    Solo se apartan los marcadores genéricos ([TABLA], [IMAGEN]...). Los de
    imagen con nombre —IMAGEN_1, IMAGEN_MAPA— se quedan como variables a
    propósito: en el editor se rellenan desde la cajita pegando la imagen
    (el valor lleva su marcador [[IMAGEN: key]]), así que si el índice los
    esconde no hay desde dónde rellenar la Ilustración N°9. Mismo criterio
    que la estructura por sección.
    """
    existing = {v["key"] for v in node["variables"]}
    for match in VAR_RE.finditer(text):
        key = match.group(1)
        upper = key.upper()
        if upper in TABLE_KEYS or upper in IMAGE_KEYS or key in existing:
            continue
        existing.add(key)
        node["variables"].append(
            {"key": key, "label": key.replace("_", " ").title()}
        )


def _collect_table_variables(table: Table, node: Dict[str, Any]) -> None:
    """Variables de todas las celdas de una tabla, colgadas del nodo dado.

    Una celda combinada aparece repetida en `row.cells`; no importa, porque
    `_collect_variables` ya descarta las claves repetidas.
    """
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _collect_variables(para.text, node)


def build_outline(docx_buffer: bytes) -> List[Dict[str, Any]]:
    """
    Árbol de títulos del documento. Cada nodo:
        level, text, para_idx, bookmark, variables[], children[]

    Las variables se cuelgan del título más profundo que esté abierto, igual
    criterio que usa la estructura por sección.
    """
    doc = Document(BytesIO(docx_buffer))

    outline: List[Dict[str, Any]] = []
    stack: List[Dict[str, Any]] = []  # títulos abiertos, por profundidad

    # Se recorre el cuerpo entero, párrafos y tablas intercalados en su orden
    # real. Recorrer `doc.paragraphs` se saltaba las tablas — sus párrafos no
    # están en esa lista — y las variables de dentro (los [CODIGO_M1],
    # [IE_ALTER1]... de los cuadros comparativos) no salían en el índice.
    #
    # `idx` cuenta solo párrafos: es la posición dentro de `doc.paragraphs`,
    # que es de donde `heading_bookmark` deriva el nombre del marcador.
    idx = -1
    for block in doc.iter_inner_content():
        if isinstance(block, Table):
            # Las variables de la tabla cuelgan del título abierto, como
            # cualquier variable del texto que la rodea
            if stack:
                _collect_table_variables(block, stack[-1])
            continue

        para = block
        idx += 1
        text = para.text
        level = _get_heading_level(para)

        if level is not None and text.strip():
            level = refine_level_with_numbering(text.strip(), level)
            clean = VAR_RE.sub("", text).replace("\n", " ").strip() or text.strip()
            node: Dict[str, Any] = {
                "level":     level,
                "text":      clean,
                "para_idx":  idx,
                "bookmark":  heading_bookmark(idx),
                "variables": [],
                "children":  [],
            }

            depth = level - 1
            stack = stack[:depth]

            if stack:
                stack[-1]["children"].append(node)
            else:
                outline.append(node)

            stack.append(node)
            # Algunas plantillas ponen la variable en el propio título
            _collect_variables(text, node)
            continue

        if stack:
            _collect_variables(text, stack[-1])

    return outline
