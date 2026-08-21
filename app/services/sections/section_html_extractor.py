# app/services/sections/section_html_extractor.py
"""
Enfoque simple y correcto:
- Cada sección se extrae directamente del .docx por rango de párrafos
- El .docx completo se cachea (evita re-descargar de MinIO)
- El HTML de cada sección se cachea (evita re-procesar con mammoth)
- Resultado: primera vez lenta, siguientes instantáneas
"""
import re
from docx import Document
from docx.table import Table
from io import BytesIO
from app.services.sections.heading_parser import (
    _get_heading_level,
    refine_level_with_numbering,
)
from app.services.html_converter import docx_to_html
from app.utils.docx_cache import get_docx_cached
from app.utils.full_html_cache import (
    get_full_html_cache,
    set_full_html_cache,
    invalidate_full_html_cache,
)

VAR_RE = re.compile(
    r'\[([A-ZÁÉÍÓÚÑÜa-záéíóúñü][A-ZÁÉÍÓÚÑÜa-záéíóúñü0-9_]*)\]',
    re.UNICODE,
)
TABLE_KEYS = {"CREAR_O_AÑADIR_TABLA","CREAR_O_ANADIR_TABLA","AÑADIR_TABLA",
              "ANADIR_TABLA","TABLA","CREAR_TABLA","INSERTAR_TABLA"}
IMAGE_KEYS = {"IMAGEN","IMAGEN_PEGAR","PEGAR_IMAGEN","INSERTAR_IMAGEN","FOTO"}


def _agregar_variables(text: str, node: dict) -> None:
    """Variables de un texto, colgadas del nodo sin repetir claves."""
    for match in VAR_RE.finditer(text):
        key = match.group(1)
        if key.upper() in TABLE_KEYS or key.upper() in IMAGE_KEYS:
            continue
        if key not in [v["key"] for v in node["variables"]]:
            node["variables"].append({
                "key":   key,
                "label": key.replace("_", " ").title(),
            })


def _agregar_variables_de_tabla(table: Table, node: dict) -> None:
    """Variables de todas las celdas de una tabla.

    Una celda combinada se repite en `row.cells`; da igual, porque las claves
    repetidas se descartan.
    """
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                _agregar_variables(para.text, node)


def _get_h1_ranges(doc: Document) -> list:
    """
    Retorna lista de { index, start, end } para cada H1.
    start y end son los paragraph_index del rango que le pertenece.
    """
    paragraphs = doc.paragraphs
    total      = len(paragraphs)

    h1_list = [
        idx for idx, para in enumerate(paragraphs)
        if _get_heading_level(para) == 1
    ]

    ranges = []
    for i, h1_idx in enumerate(h1_list):
        end = h1_list[i + 1] if i + 1 < len(h1_list) else total
        ranges.append({
            "index": h1_idx,
            "start": h1_idx,
            "end":   end,
        })

    return ranges


def _extract_docx_section(docx_buffer: bytes, start: int, end: int) -> bytes:
    """
    Crea un nuevo .docx con solo los párrafos del rango [start, end).
    Preserva estilos y formato del original.
    """
    new_doc       = Document(BytesIO(docx_buffer))
    body          = new_doc.element.body
    para_idx      = 0
    last_included = False
    to_remove     = []

    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if not (start <= para_idx < end):
                to_remove.append(child)
                last_included = False
            else:
                last_included = True
            para_idx += 1

        elif tag == "tbl":
            if not last_included:
                to_remove.append(child)

    for el in to_remove:
        try:
            body.remove(el)
        except Exception:
            pass

    output = BytesIO()
    new_doc.save(output)
    return output.getvalue()


def _extract_structure(doc: Document, start: int, end: int) -> list:
    """
    Construye árbol de headings + variables.
    heading_stack[i] es el nodo activo en profundidad i (0=H2, 1=H3, 2=H4…).
    Las variables se asignan al nodo MÁS PROFUNDO activo, no siempre al H2.
    """
    structure     = []
    heading_stack: list = []  # stack de nodos activos por profundidad

    # Cuerpo entero en su orden real, tablas incluidas: `doc.paragraphs` no
    # contiene los párrafos de dentro de una tabla, y con el rango solo de
    # párrafos las variables de los cuadros ([CODIGO_M1], [IE_ALTER1]...)
    # no aparecían en la estructura de la sección.
    idx = -1
    for block in doc.iter_inner_content():
        if isinstance(block, Table):
            # La tabla pertenece a la sección si el último párrafo visto cae
            # dentro del rango — mismo criterio que usa _extract_docx_section
            if start <= idx < end and heading_stack:
                _agregar_variables_de_tabla(block, heading_stack[-1])
            continue

        para = block
        idx += 1
        if idx < start:
            continue
        if idx >= end:
            break
        level = _get_heading_level(para)
        text  = para.text.strip()
        if level is not None and text:
            level = refine_level_with_numbering(text, level)

        if level == 1 and idx == start:
            # Crear nodo raíz para variables que aparecen directamente bajo el H1
            # (capítulos sin subsecciones H2, ej: Capítulo 5 con una sola variable)
            clean_h1 = VAR_RE.sub('', text).replace('\n', ' ').strip()
            h1_root = {
                "level": 2, "text": clean_h1 or text,
                "para_idx": idx, "variables": [], "children": [],
            }
            heading_stack = [h1_root]
            # Se agrega a structure solo si termina con variables
            continue

        if level is not None and level >= 2:
            # Limpiar el texto del heading: quitar placeholders [VAR] y saltos blandos
            clean_text = VAR_RE.sub('', text).replace('\n', ' ').strip()
            node = {
                "level": level, "text": clean_text or text,
                "para_idx": idx, "variables": [], "children": [],
            }
            depth = level - 2  # 0 → H2, 1 → H3, 2 → H4 …

            # Si el h1_root (stack[0]) tiene variables, publicarlo antes de añadir el H2
            if depth == 0 and len(heading_stack) == 1 and heading_stack[0].get("variables"):
                structure.append(heading_stack[0])

            # Recortar el stack a la profundidad del nodo padre
            heading_stack = heading_stack[:depth]

            if depth == 0:
                structure.append(node)
            elif heading_stack:
                heading_stack[-1]["children"].append(node)
            else:
                structure.append(node)

            heading_stack.append(node)

            # Algunos templates ponen [VAR] dentro del mismo párrafo del heading.
            # Usar para.text (recorre TODOS los <w:t> recursivamente, incluyendo
            # texto dentro de <w:hyperlink> o <w:ins> que para.runs no ve).
            _agregar_variables(para.text, node)

        else:
            # Párrafo de cuerpo — asignar variables al nodo más profundo activo
            if not heading_stack:
                continue
            # para.text en lugar de para.runs para capturar texto en hyperlinks/ins
            _agregar_variables(para.text, heading_stack[-1])

    # Publicar h1_root si tiene variables directas o hijos (H3 sin H2 intermedios)
    # y no fue ya publicado. El check len==1 fallaba cuando había H3 en el stack.
    if heading_stack and heading_stack[0] not in structure:
        root = heading_stack[0]
        if root.get("variables") or root.get("children"):
            structure.insert(0, root)

    return structure


def _build_full_cache(minio_key: str) -> dict:  # alias público para warmup
    return _build_cache(minio_key)


def _build_cache(minio_key: str) -> dict:
    """
    Procesa TODAS las secciones del documento y las cachea.
    Se llama una sola vez por documento.
    """
    docx_buffer = get_docx_cached(minio_key)
    doc         = Document(BytesIO(docx_buffer))
    h1_ranges   = _get_h1_ranges(doc)

    sections_html      = {}
    sections_structure = {}

    for h1 in h1_ranges:
        h1_idx  = h1["index"]
        start   = h1["start"]
        end     = h1["end"]

        # Extraer .docx de la sección y convertir a HTML
        section_buffer          = _extract_docx_section(docx_buffer, start, end)
        html_result             = docx_to_html(section_buffer)
        sections_html[h1_idx]      = html_result["html"]
        sections_structure[h1_idx] = _extract_structure(doc, start, end)

    cache_data = {
        "sections":  sections_html,
        "structure": sections_structure,
    }

    set_full_html_cache(minio_key, cache_data)
    return cache_data


def get_section_full(minio_key: str, h1_index: int) -> dict:
    """
    Retorna { html, structure } para la sección indicada.
    Si el caché no existe lo construye completo.
    Si la sección específica no está (raro), la procesa al vuelo.
    """
    cached = get_full_html_cache(minio_key)

    # Caché no existe → construir todo
    if cached is None:
        cached = _build_cache(minio_key)

    html      = cached["sections"].get(h1_index)
    structure = cached["structure"].get(h1_index, [])

    # Sección no encontrada en caché → procesar solo esa
    if html is None:
        docx_buffer = get_docx_cached(minio_key)
        doc         = Document(BytesIO(docx_buffer))
        h1_ranges   = _get_h1_ranges(doc)

        target = next((r for r in h1_ranges if r["index"] == h1_index), None)
        if target is None:
            raise ValueError(
                f"h1_index={h1_index} no existe en el documento."
            )

        section_buffer = _extract_docx_section(
            docx_buffer, target["start"], target["end"]
        )
        html_result = docx_to_html(section_buffer)
        html        = html_result["html"]
        structure   = _extract_structure(doc, target["start"], target["end"])

        # Actualizar caché
        cached["sections"][h1_index]  = html
        cached["structure"][h1_index] = structure
        set_full_html_cache(minio_key, cached)

    return {
        "html":      html,
        "messages":  [],
        "structure": structure,
    }


def extract_section_as_html(minio_key: str, h1_index: int) -> dict:
    result = get_section_full(minio_key, h1_index)
    return {"html": result["html"], "messages": result["messages"]}


def extract_section_structure(minio_key: str, h1_index: int) -> list:
    return get_section_full(minio_key, h1_index)["structure"]


def invalidate_document_cache(minio_key: str) -> None:
    from app.utils.docx_cache import invalidate_cache
    invalidate_cache(minio_key)
    invalidate_full_html_cache(minio_key)