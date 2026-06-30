# app/services/sections/section_extractor.py
import copy
from docx import Document
from docx.oxml.ns import qn
from io import BytesIO
from typing import List, Dict, Set, Tuple
from app.services.sections.heading_parser import _get_heading_level


def _build_section_ranges(doc: Document) -> Dict[int, Tuple[int, int]]:
    """Mapea h1_para_index → (start_child_idx, end_child_idx) en el body XML."""
    paragraphs = doc.paragraphs
    total_paras = len(paragraphs)
    h1_indexes = [
        idx for idx, para in enumerate(paragraphs)
        if _get_heading_level(para) == 1
    ]

    body = doc.element.body
    children = list(body)

    para_to_child: Dict[int, int] = {}
    p_idx = 0
    for c_idx, child in enumerate(children):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            para_to_child[p_idx] = c_idx
            p_idx += 1

    total_children = len(children)
    ranges: Dict[int, Tuple[int, int]] = {}

    for i, h1_idx in enumerate(h1_indexes):
        next_h1 = h1_indexes[i + 1] if i + 1 < len(h1_indexes) else total_paras
        start_child = para_to_child.get(h1_idx, 0)
        end_child = para_to_child.get(next_h1, total_children)
        ranges[h1_idx] = (start_child, end_child)

    return ranges


def extract_sections(docx_buffer: bytes, selected_para_indexes: List[int]) -> bytes:
    """
    Extrae secciones en el ORDEN que indica selected_para_indexes.
    Clona el doc original para preservar estilos, luego reordena el body
    según la secuencia solicitada.
    """
    doc = Document(BytesIO(docx_buffer))
    section_ranges = _build_section_ranges(doc)

    valid = [idx for idx in selected_para_indexes if idx in section_ranges]
    if not valid:
        raise ValueError("Ninguna sección encontrada con los índices proporcionados")

    new_doc = Document(BytesIO(docx_buffer))
    body = new_doc.element.body
    original_children = list(body)

    sect_pr = None
    for child in original_children:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "sectPr":
            sect_pr = child

    # Encontrar dónde empieza el primer H1 → todo lo anterior es preámbulo (TOC, portada)
    first_h1_child = min(start for start, _ in section_ranges.values()) if section_ranges else 0

    for child in list(body):
        body.remove(child)

    # Preámbulo: portada, tabla de índices, etc.
    for c_idx in range(first_h1_child):
        body.append(copy.deepcopy(original_children[c_idx]))

    # Secciones en el orden elegido por el usuario
    for h1_idx in valid:
        start, end = section_ranges[h1_idx]
        for c_idx in range(start, end):
            if c_idx < len(original_children):
                body.append(copy.deepcopy(original_children[c_idx]))

    if sect_pr is not None:
        body.append(copy.deepcopy(sect_pr))

    output = BytesIO()
    new_doc.save(output)
    return output.getvalue()