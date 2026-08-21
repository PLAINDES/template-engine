# app/services/filler.py
import re
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from typing import Dict, List

VAR_RE = re.compile(
    r'\[([A-ZÁÉÍÓÚÑÜa-záéíóúñü][A-ZÁÉÍÓÚÑÜa-záéíóúñü0-9_]*)\]',
    re.UNICODE,
)

TABLE_KEYS = {
    "CREAR_O_AÑADIR_TABLA", "CREAR_O_ANADIR_TABLA",
    "AÑADIR_TABLA", "ANADIR_TABLA",
    "TABLA", "CREAR_TABLA", "INSERTAR_TABLA",
}

IMAGE_KEYS = {
    "IMAGEN", "IMAGEN_PEGAR", "PEGAR_IMAGEN",
    "INSERTAR_IMAGEN", "FOTO", "ANADIR_IMAGE",
}

IMAGE_PREFIXES = ("IMAGEN_", "FOTO_", "IMG_")


def _is_image_key(upper_key: str) -> bool:
    if upper_key in IMAGE_KEYS:
        return True
    for prefix in IMAGE_PREFIXES:
        if upper_key.startswith(prefix) and upper_key not in IMAGE_KEYS:
            return True
    return False


def fill_document(
    docx_buffer: bytes,
    variables:   Dict[str, str],
    tablas:      List[Dict],
    imagenes:    List[Dict],
    bloques:     List[Dict] = [],
    replacements: List[Dict] = [],
) -> bytes:
    print(f"[filler] tablas recibidas: {len(tablas)}")
    for t in tablas:
        print(f"  tabla: para_idx={t.get('paragraph_index')} headers={t.get('headers')}")
    print(f"[filler] imagenes recibidas: {len(imagenes)}")
    for i in imagenes:
        print(f"  imagen: para_idx={i.get('paragraph_index')} key={i.get('minio_key')}")
    print(f"[filler] bloques recibidos: {len(bloques)}")
    for b in bloques:
        print(f"  bloque tipo={b.get('tipo')} items={len(b.get('items', []))}")
    doc = Document(BytesIO(docx_buffer))

    # 0. Bloques repetibles (antes de variables para que las replique también)
    if bloques:
        _expand_blocks(doc, bloques)

    # 1. Variables de texto
    for para in doc.paragraphs:
        _replace_vars_in_paragraph(para, variables, doc)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_vars_in_paragraph(para, variables, doc)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_vars_in_paragraph(para, variables)
        for para in section.footer.paragraphs:
            _replace_vars_in_paragraph(para, variables)

    # 2. Reemplazos de texto (mejoras de IA)
    if replacements:
        print(f"[filler] reemplazos de texto: {len(replacements)}")
        for repl in replacements:
            orig = repl.get("originalText", "")
            nuevo = repl.get("newText", "")
            if orig and nuevo:
                _apply_text_replacement(doc, orig, nuevo)

    # 3. Tablas
    tablas_sorted = sorted(tablas, key=lambda t: t.get("paragraph_index", 0), reverse=True)
    for tabla_data in tablas_sorted:
        _insert_table_at_placeholder(doc, tabla_data)

    # 3. Imágenes
    imagenes_sorted = sorted(imagenes, key=lambda i: i.get("paragraph_index", 0), reverse=True)
    for imagen_data in imagenes_sorted:
        _insert_image_at_placeholder(doc, imagen_data)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


_ANNOTATION_RE = re.compile(
    r'\[(?:AÑADIR|ANADIR|INSERTAR|AGREGAR|AÑADE|AÑADA)_?(?:IMAGEN|CUADRO|TABLA|MAPA|GRAFICO|GRÁFICO)[^\]]*\]',
    re.IGNORECASE | re.UNICODE,
)


# Bloque de tabla dentro del texto de una variable. La IA lo emite así:
#   [[TABLA: Título de la tabla]]
#   Encabezado | Encabezado | Encabezado
#   celda | celda | celda
#   [[/TABLA]]
# y aquí se convierte en una tabla real de Word — pegado como texto plano
# quedaba un choclo ilegible de barras justificadas.
_TABLA_START_RE = re.compile(r'^\[\[TABLA(?::\s*(.*?))?\]\]$')
_TABLA_END = '[[/TABLA]]'

# Imagen dentro del texto de una variable. La escribe el editor cuando el
# usuario pega, arrastra o elige una imagen en la cajita de la variable:
#   [[IMAGEN: docx-images/1717000000_mapa.png]]
#   [[IMAGEN: docx-images/1717000000_mapa.png | Fuente: elaboración propia]]
# La key es la de MinIO que devolvió /upload-image.
_IMAGEN_INLINE_RE = re.compile(
    r'^\[\[IMAGEN:\s*([^|\]]+?)\s*(?:\|\s*([^\]]*?)\s*)?\]\]$'
)

# Ancho de la imagen insertada en el texto. Es el mismo que usa el modal de
# imágenes con nombre, para que las dos vías produzcan la misma página.
ANCHO_IMAGEN_EN_LINEA = 5.0


def _partir_en_segmentos(lines: List[str]):
    """Separa las líneas del valor en texto corrido, bloques de tabla e imágenes."""
    segmentos = []  # ("texto", línea) | ("tabla", título, filas) | ("imagen", key, pie)
    i = 0
    while i < len(lines):
        img = _IMAGEN_INLINE_RE.match(lines[i].strip())
        if img:
            segmentos.append(("imagen", img.group(1), (img.group(2) or '').strip()))
            i += 1
            continue
        m = _TABLA_START_RE.match(lines[i].strip())
        if m:
            titulo = (m.group(1) or '').strip()
            filas = []
            i += 1
            while i < len(lines) and lines[i].strip() != _TABLA_END:
                filas.append([c.strip() for c in lines[i].split('|')])
                i += 1
            i += 1  # saltar el [[/TABLA]]
            if filas:
                segmentos.append(("tabla", titulo, filas))
        else:
            segmentos.append(("texto", lines[i]))
            i += 1
    return segmentos


def _poner_bordes_de_tabla(table) -> None:
    """Bordes sencillos en toda la tabla, dibujados directo en el XML."""
    tbl_pr = table._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borde = OxmlElement(f"w:{lado}")
        borde.set(qn("w:val"), "single")
        borde.set(qn("w:sz"), "4")
        borde.set(qn("w:color"), "000000")
        borders.append(borde)
    tbl_pr.append(borders)


def _construir_tabla(doc: Document, filas: List[List[str]]):
    """Tabla real con la primera fila como encabezado en negrita."""
    num_cols = max(len(f) for f in filas)
    table = doc.add_table(rows=len(filas), cols=num_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        # La plantilla no define ese estilo (los .docx solo llevan los estilos
        # que usan): sin esto la tabla saldría invisible, así que los bordes
        # se dibujan directamente
        _poner_bordes_de_tabla(table)
    for ri, fila in enumerate(filas):
        for ci, valor in enumerate(fila):
            if ci >= num_cols:
                continue
            cell = table.rows[ri].cells[ci]
            cell.text = str(valor)
            if ri == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    return table


def _insertar_imagen_en_linea(doc: Document, anchor, minio_key: str, pie: str):
    """
    Coloca la imagen que el usuario pegó dentro del valor de una variable.

    Devuelve el último elemento insertado para que quien recorre los segmentos
    siga encadenando detrás de él y el orden del texto se conserve.
    """
    from app.utils.minio_client import download_from_minio

    try:
        image_bytes = download_from_minio(minio_key)
    except Exception as e:
        print(f"[filler] error descargando imagen en linea [{minio_key}]: {e}")
        return anchor

    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.add_run().add_picture(BytesIO(image_bytes), width=Inches(ANCHO_IMAGEN_EN_LINEA))
    anchor.addnext(img_para._element)
    anchor = img_para._element

    if pie:
        pie_para = doc.add_paragraph()
        run = pie_para.add_run(pie)
        run.font.size      = Pt(9)
        run.font.italic    = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        pie_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        anchor.addnext(pie_para._element)
        anchor = pie_para._element

    return anchor


def _replace_vars_in_paragraph(para, variables: Dict[str, str], doc: Document = None):
    if not para.runs:
        return
    full_text = "".join(run.text for run in para.runs)
    if not VAR_RE.search(full_text):
        return

    new_text = full_text
    for key, value in variables.items():
        # Limpiar anotaciones de imagen/cuadro que la IA genera dentro del texto
        clean_value = _ANNOTATION_RE.sub('', str(value)).strip()
        new_text = new_text.replace(f"[{key}]", clean_value)

    if new_text == full_text:
        return

    # Si el valor tiene saltos de línea, expandir en múltiples párrafos
    lines = [ln for ln in new_text.split('\n') if ln.strip()]
    if not lines:
        # Todo era anotaciones — dejar el párrafo vacío
        if para.runs:
            para.runs[0].text = ''
            for run in para.runs[1:]:
                run.text = ''
        return

    # Sin documento (cabeceras y pies) los bloques de tabla van como texto y las
    # imagenes no se pueden incrustar: se descarta el marcador en vez de dejarlo a la vista
    segmentos = (
        _partir_en_segmentos(lines) if doc is not None
        else [("texto", ln) for ln in lines if not _IMAGEN_INLINE_RE.match(ln.strip())]
    )

    # El primer segmento de texto va en este mismo párrafo (conserva su estilo)
    primero = segmentos[0] if segmentos else None
    if primero is not None and primero[0] == "texto":
        para.runs[0].text = primero[1]
        for run in para.runs[1:]:
            run.text = ''
        segmentos = segmentos[1:]
    else:
        # El valor arranca con una tabla: el placeholder queda vacío
        for run in para.runs:
            run.text = ''

    # El resto se inserta en orden después del placeholder
    anchor = para._element
    for seg in segmentos:
        if seg[0] == "texto":
            new_para = copy.deepcopy(para)
            if new_para.runs:
                new_para.runs[0].text = seg[1]
                for run in new_para.runs[1:]:
                    run.text = ''
            anchor.addnext(new_para._element)
            anchor = new_para._element
        elif seg[0] == "imagen":
            _, minio_key, pie = seg
            anchor = _insertar_imagen_en_linea(doc, anchor, minio_key, pie)
        else:
            _, titulo, filas = seg
            if titulo:
                titulo_para = doc.add_paragraph()
                run = titulo_para.add_run(titulo)
                run.font.size = Pt(10)
                run.font.bold = True
                titulo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                anchor.addnext(titulo_para._element)
                anchor = titulo_para._element
            table = _construir_tabla(doc, filas)
            anchor.addnext(table._element)
            anchor = table._element


def _add_caption_paragraph(doc: Document, target_para, text: str, align: str = "center") -> None:
    """Inserta un párrafo de caption (título o pie) junto al target_para."""
    if not text or not text.strip():
        return
    caption_para = doc.add_paragraph()
    run = caption_para.add_run(text.strip())
    run.font.size    = Pt(9)
    run.font.italic  = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    alignment_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    caption_para.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    # Mover el párrafo al lugar correcto en el XML
    target_para._element.addnext(caption_para._element)


def _insert_table_at_placeholder(doc: Document, tabla_data: Dict):
    para_idx      = tabla_data.get("paragraph_index", 0)
    headers       = tabla_data.get("headers", [])
    rows_data     = tabla_data.get("rows", [])
    titulo        = tabla_data.get("titulo", "")
    pie           = tabla_data.get("pie", "")

    if not headers and not rows_data:
        return

    paragraphs = doc.paragraphs
    print(f"[insert_table] buscando párrafo {para_idx}")
    print(f"[insert_table] total párrafos en doc: {len(paragraphs)}")
    if para_idx < len(paragraphs):
        text = "".join(run.text for run in paragraphs[para_idx].runs)
        print(f"[insert_table] texto en párrafo {para_idx}: '{text}'")
    if para_idx >= len(paragraphs):
        return

    target_para = paragraphs[para_idx]
    full_text   = "".join(run.text for run in target_para.runs)
    has_placeholder = any(k in full_text for k in TABLE_KEYS)

    if not has_placeholder:
        for para in paragraphs:
            t = "".join(run.text for run in para.runs)
            if any(k in t for k in TABLE_KEYS):
                target_para = para
                break
        else:
            return

    # Limpiar placeholder
    for run in target_para.runs:
        run.text = ""

    num_cols = max(len(headers), max((len(r) for r in rows_data), default=0))
    if num_cols == 0:
        return

    num_rows = len(rows_data) + (1 if headers else 0)
    table    = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"

    row_offset = 0
    if headers:
        header_row = table.rows[0]
        for ci, header_text in enumerate(headers):
            cell = header_row.cells[ci]
            cell.text = str(header_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        row_offset = 1

    for ri, row_data in enumerate(rows_data):
        table_row = table.rows[ri + row_offset]
        for ci, cell_val in enumerate(row_data):
            if ci < num_cols:
                table_row.cells[ci].text = str(cell_val)

    # Insertar tabla en el XML
    target_para._element.addnext(table._element)

    # Pie debajo de la tabla
    if pie and pie.strip():
        pie_para = doc.add_paragraph()
        run = pie_para.add_run(pie.strip())
        run.font.size    = Pt(9)
        run.font.italic  = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        pie_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table._element.addnext(pie_para._element)

    # Título encima de la tabla
    if titulo and titulo.strip():
        titulo_para = doc.add_paragraph()
        run = titulo_para.add_run(titulo.strip())
        run.font.size   = Pt(10)
        run.font.bold   = True
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        target_para._element.addnext(titulo_para._element)


def _insert_image_at_placeholder(doc: Document, imagen_data: Dict):
    from app.utils.minio_client import download_from_minio

    para_idx     = imagen_data.get("paragraph_index", 0)
    minio_key    = imagen_data.get("minio_key", "")
    var_key      = imagen_data.get("key", "")
    width_inches = imagen_data.get("width_inches", 4.0)
    titulo       = imagen_data.get("titulo", "")
    pie          = imagen_data.get("pie", "")
    descripcion  = imagen_data.get("descripcion", "")

    if not minio_key:
        return

    paragraphs  = doc.paragraphs
    target_para = None

    # Match por nombre de variable (imágenes con nombre único)
    if var_key:
        placeholder = f"[{var_key}]"
        for para in paragraphs:
            t = "".join(run.text for run in para.runs)
            if placeholder in t:
                target_para = para
                break

    # Fallback: match por paragraph_index (imágenes genéricas)
    if not target_para and para_idx < len(paragraphs):
        t = "".join(run.text for run in paragraphs[para_idx].runs)
        if any(k in t.upper() for k in IMAGE_KEYS) or any(t.upper().startswith(p) for p in IMAGE_PREFIXES):
            target_para = paragraphs[para_idx]

    # Fallback: buscar cualquier placeholder de imagen
    if not target_para:
        for para in paragraphs:
            t = "".join(run.text for run in para.runs)
            if _is_image_key(t.strip().strip("[]").upper()):
                target_para = para
                break

    if not target_para:
        return

    # Limpiar placeholder
    for run in target_para.runs:
        run.text = ""

    try:
        image_bytes  = download_from_minio(minio_key)
        image_buffer = BytesIO(image_bytes)
    except Exception as e:
        print(f"Error descargando imagen [{minio_key}]: {e}")
        return

    # Insertar imagen
    run = target_para.add_run()
    run.add_picture(image_buffer, width=Inches(width_inches))
    target_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Pie debajo de la imagen
    if pie and pie.strip():
        pie_para = doc.add_paragraph()
        run = pie_para.add_run(pie.strip())
        run.font.size      = Pt(9)
        run.font.italic    = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        pie_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        target_para._element.addnext(pie_para._element)

    # Descripción debajo de la imagen (párrafo normal, texto generado por IA)
    if descripcion and descripcion.strip():
        desc_para = doc.add_paragraph()
        run = desc_para.add_run(descripcion.strip())
        run.font.size = Pt(11)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Insertar después del pie si existe, si no después de la imagen
        anchor = target_para._element
        if pie and pie.strip():
            # Buscar el pie que acabamos de insertar
            next_elem = target_para._element.getnext()
            if next_elem is not None:
                anchor = next_elem
        anchor.addnext(desc_para._element)

    # Título encima de la imagen — addprevious para que quede antes, no después
    if titulo and titulo.strip():
        titulo_para = doc.add_paragraph()
        run = titulo_para.add_run(titulo.strip())
        run.font.size  = Pt(10)
        run.font.bold  = True
        titulo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        target_para._element.addprevious(titulo_para._element)


def _apply_text_replacement(doc: Document, original: str, replacement: str) -> None:
    """Busca el texto original en los párrafos del documento y lo reemplaza."""
    replaced = False
    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        if original in full_text:
            new_text = full_text.replace(original, replacement)
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            replaced = True
            print(f"[filler] ✓ Reemplazo aplicado ({len(original)} → {len(replacement)} chars)")
            break
    if not replaced:
        # Buscar en celdas de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        full_text = "".join(run.text for run in para.runs)
                        if original in full_text:
                            new_text = full_text.replace(original, replacement)
                            if para.runs:
                                para.runs[0].text = new_text
                                for run in para.runs[1:]:
                                    run.text = ""
                            replaced = True
                            print(f"[filler] ✓ Reemplazo en tabla ({len(original)} → {len(replacement)} chars)")
                            break
                    if replaced:
                        break
                if replaced:
                    break
            if replaced:
                break
    if not replaced:
        print(f"[filler] ✗ No se encontró texto para reemplazar: '{original[:50]}...'")


def _find_preceding_heading_elem(body, start_elem):
    """
    Busca el párrafo con estilo Heading-1 más cercano que precede a start_elem.
    Retorna una copia profunda del elemento XML, o None si no encuentra.
    Detecta estilos en cualquier idioma: Heading1, Ttulo1, h1, H1, etc.
    """
    preceding = []
    for elem in list(body):
        if elem is start_elem:
            break
        preceding.append(elem)

    for elem in reversed(preceding):
        if not elem.tag.endswith('}p'):
            continue
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            continue
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            continue
        style_val = pStyle.get(qn('w:val'), '') or ''
        # Captura Heading1, Ttulo1 (ES), berschrift1 (DE), h1, H1, etc.
        if re.search(r'(?i)[a-z]1$', style_val):
            return copy.deepcopy(elem)
    return None


def _expand_blocks(doc: Document, bloques: List[Dict]) -> None:
    """
    Detecta TODOS los pares [BLOQUE_X_START]/[BLOQUE_X_END] en el documento
    en un único escaneo (antes de modificar nada), luego expande cada par.

    Para bloques con más de 1 ítem, replica automáticamente el heading H1 que
    precede al bloque al inicio de cada repetición (a partir de la 2ª), de modo
    que cada empresa/ítem quede bajo su propio encabezado de sección.
    """
    body = doc.element.body

    for bloque in bloques:
        tipo  = bloque.get("tipo", "").upper()
        items = bloque.get("items", [])
        if not items:
            continue

        start_marker = f"[BLOQUE_{tipo}_START]"
        end_marker   = f"[BLOQUE_{tipo}_END]"

        # ── Paso 1: encontrar TODOS los pares usando referencias a elementos ──
        pairs = []          # lista de (start_elem, end_elem, [template_elems])
        current_start = None

        for child in list(body):
            if not child.tag.endswith('}p'):
                continue
            text = "".join(t.text or "" for t in child.iter() if t.tag.endswith('}t'))
            if start_marker in text and current_start is None:
                current_start = child
            elif end_marker in text and current_start is not None:
                # Recolectar los elementos entre start y end
                collecting = False
                template_elems = []
                for elem in list(body):
                    if elem is current_start:
                        collecting = True
                        continue
                    if elem is child:
                        break
                    if collecting:
                        template_elems.append(copy.deepcopy(elem))
                pairs.append((current_start, child, template_elems))
                current_start = None

        if not pairs:
            print(f"[expand_blocks] No se encontraron pares para tipo={tipo}")
            continue

        print(f"[expand_blocks] tipo={tipo} pares={len(pairs)} items={len(items)}")

        # Heading H1 que precede al bloque — se replica para ítems 2, 3, …
        preceding_heading = (
            _find_preceding_heading_elem(body, pairs[0][0])
            if len(items) > 1
            else None
        )
        if preceding_heading is not None:
            heading_preview = "".join(
                t.text or "" for t in preceding_heading.iter() if t.tag.endswith('}t')
            )
            print(f"[expand_blocks] heading previo detectado: '{heading_preview[:80]}'")

        # ── Paso 2: expandir cada par (los pares no se solapan, orden seguro) ──
        for start_elem, end_elem, template_elements in pairs:
            for item_idx, item in enumerate(reversed(items)):
                is_first_item = (item_idx == len(items) - 1)

                # Insertar contenido del bloque en orden correcto (reversed + addnext)
                for elem in reversed(template_elements):
                    new_elem = copy.deepcopy(elem)
                    _replace_vars_in_xml_element(new_elem, item)
                    end_elem.addnext(new_elem)

                # Para ítems 2+: insertar copia del heading justo antes del bloque
                if not is_first_item and preceding_heading is not None:
                    heading_copy = copy.deepcopy(preceding_heading)
                    _replace_vars_in_xml_element(heading_copy, item)
                    end_elem.addnext(heading_copy)

            # Eliminar desde start hasta end inclusive
            to_remove = []
            collecting = False
            for child in list(body):
                if child is start_elem:
                    collecting = True
                if collecting:
                    to_remove.append(child)
                if child is end_elem:
                    break
            for elem in to_remove:
                try:
                    body.remove(elem)
                except Exception:
                    pass


def _replace_vars_in_xml_element(element, variables: Dict[str, str]) -> None:
    """Reemplaza variables [KEY] en nodos <w:t> del elemento XML."""
    for node in element.iter():
        # Solo <w:t> tiene text editable; CT_P y otros tienen text como property read-only
        if node.tag.endswith('}t') and node.text:
            for key, value in variables.items():
                node.text = node.text.replace(f"[{key}]", str(value))