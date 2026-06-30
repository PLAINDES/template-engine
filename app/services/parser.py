# app/services/parser.py
import re
from docx import Document
from io import BytesIO
from typing import Dict, List

# Regex flexible:
# - acepta MAYÚSCULAS y minúsculas
# - acepta acentos y Ñ
# - acepta números y guiones bajos
# - el primer carácter debe ser letra
#
# Ejemplos válidos:
# [NOMBRE]
# [nombre]
# [Nombre_1]
# [brecha_1]
VAR_RE = re.compile(
    r'\[([A-ZÁÉÍÓÚÑÜa-záéíóúñü][A-ZÁÉÍÓÚÑÜa-záéíóúñü0-9_]*)\]',
    re.UNICODE,
)

# Keywords de tabla — todas las variantes posibles
TABLE_KEYS = {
    "CREAR_O_AÑADIR_TABLA",
    "CREAR_O_ANADIR_TABLA",
    "AÑADIR_TABLA",
    "ANADIR_TABLA",
    "TABLA",
    "CREAR_TABLA",
    "INSERTAR_TABLA",
}

# Keywords de imagen — variantes genéricas (sin nombre único)
IMAGE_KEYS = {
    "IMAGEN",
    "IMAGEN_PEGAR",
    "PEGAR_IMAGEN",
    "INSERTAR_IMAGEN",
    "FOTO",
    "FOTOGRAFÍA",
    "FOTOGRAFIA",
    "ANADIR_IMAGE",
}

# Prefijos que indican una imagen con nombre único: [IMAGEN_MAPA_UBICACION]
IMAGE_PREFIXES = ("IMAGEN_", "FOTO_", "IMG_")


def _is_image_key(upper_key: str) -> bool:
    """Detecta si una variable es un placeholder de imagen (genérico o con nombre)."""
    if upper_key in IMAGE_KEYS:
        return True
    # Variables con nombre: IMAGEN_DIAGRAMA, FOTO_PLANO, IMG_MAPA
    # Pero excluir las que son keywords genéricas como IMAGEN_PEGAR
    for prefix in IMAGE_PREFIXES:
        if upper_key.startswith(prefix) and upper_key not in IMAGE_KEYS:
            return True
    return False


def extract_variables(docx_buffer: bytes) -> Dict:
    """
    Parsea un .docx y extrae variables, tablas e imágenes
    EN EL ORDEN EXACTO en que aparecen en el documento.

    IMPORTANTE:
    Word fragmenta el texto en múltiples runs.

    Ejemplo:
        run1 -> "[NOMBRE"
        run2 -> "_CLIENTE]"

    Por eso concatenamos todos los runs antes
    de aplicar regex.
    """

    doc = Document(BytesIO(docx_buffer))

    all_elements: List[Dict] = []
    seen_vars = set()

    global_idx = 0
    tabla_count = 0
    img_count = 0

    # ─────────────────────────────────────────────────────────────
    # 1. PÁRRAFOS DEL BODY
    # ─────────────────────────────────────────────────────────────
    for para_idx, para in enumerate(doc.paragraphs):

        full_text = _join_runs(para)

        if not full_text.strip():
            continue

        for match in VAR_RE.finditer(full_text):

            key = match.group(1)

            # Normalizamos para comparar keywords especiales
            upper_key = key.upper()

            # ── TABLAS ───────────────────────────────────────────
            if upper_key in TABLE_KEYS:

                all_elements.append({
                    "type": "tabla",
                    "data": {
                        "index": tabla_count,
                        "paragraph_index": para_idx,
                        "order": global_idx,
                    }
                })

                tabla_count += 1
                global_idx += 1
                continue

            # ── IMÁGENES (genéricas y con nombre) ──────────────
            if _is_image_key(upper_key):

                all_elements.append({
                    "type": "imagen",
                    "data": {
                        "index": img_count,
                        "paragraph_index": para_idx,
                        "order": global_idx,
                        "key": key,
                    }
                })

                img_count += 1
                global_idx += 1
                continue

            # ── VARIABLES NORMALES ──────────────────────────────
            if key not in seen_vars:

                seen_vars.add(key)

                all_elements.append({
                    "type": "var",
                    "data": {
                        "key": key,
                        "label": _key_to_label(key),
                        "value": "",
                        "in_table": False,
                        "order": global_idx,
                    }
                })

                global_idx += 1

    # ─────────────────────────────────────────────────────────────
    # 2. TABLAS EXISTENTES DEL DOCUMENTO
    # ─────────────────────────────────────────────────────────────
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                # Concatenar todos los párrafos de la celda
                cell_text = "\n".join(
                    _join_runs(para)
                    for para in cell.paragraphs
                )

                for match in VAR_RE.finditer(cell_text):

                    key = match.group(1)
                    upper_key = key.upper()

                    if upper_key in TABLE_KEYS or _is_image_key(upper_key):
                        continue

                    if key not in seen_vars:

                        seen_vars.add(key)

                        all_elements.append({
                            "type": "var",
                            "data": {
                                "key": key,
                                "label": _key_to_label(key),
                                "value": "",
                                "in_table": True,
                                "order": global_idx,
                            }
                        })

                        global_idx += 1

    # ─────────────────────────────────────────────────────────────
    # 3. HEADERS Y FOOTERS
    # ─────────────────────────────────────────────────────────────
    for section in doc.sections:

        header_paras = list(section.header.paragraphs)
        footer_paras = list(section.footer.paragraphs)

        for para in header_paras + footer_paras:

            full_text = _join_runs(para)

            for match in VAR_RE.finditer(full_text):

                key = match.group(1)
                upper_key = key.upper()

                if upper_key in TABLE_KEYS or _is_image_key(upper_key):
                    continue

                if key not in seen_vars:

                    seen_vars.add(key)

                    all_elements.append({
                        "type": "var",
                        "data": {
                            "key": key,
                            "label": _key_to_label(key),
                            "value": "",
                            "in_table": False,
                            "order": global_idx,
                        }
                    })

                    global_idx += 1

    # ─────────────────────────────────────────────────────────────
    # Separar por tipo manteniendo el orden
    # ─────────────────────────────────────────────────────────────
    variables = [
        e["data"]
        for e in all_elements
        if e["type"] == "var"
    ]

    tablas = [
        e["data"]
        for e in all_elements
        if e["type"] == "tabla"
    ]

    imagenes = [
        e["data"]
        for e in all_elements
        if e["type"] == "imagen"
    ]

    return {
        "variables": variables,
        "tablas": tablas,
        "imagenes": imagenes,
        "total_variables": len(variables),
        "total_tablas": len(tablas),
        "total_imagenes": len(imagenes),
    }


def debug_paragraphs(docx_buffer: bytes) -> List[Dict]:
    """
    Devuelve información de debug de párrafos
    para ver cómo Word divide los runs.
    """

    doc = Document(BytesIO(docx_buffer))
    result = []

    for i, para in enumerate(doc.paragraphs):

        full = _join_runs(para)
        runs = [r.text for r in para.runs]

        if (
            "[" in full
            or "TABLA" in full.upper()
            or "IMAGEN" in full.upper()
        ):

            result.append({
                "paragraph_index": i,
                "full_text": full,
                "runs": runs,
                "runs_count": len(runs),
            })

    return result


def _join_runs(para) -> str:
    """
    Concatena todos los runs de un párrafo
    en un solo string.
    """
    return "".join(run.text for run in para.runs)


def _key_to_label(key: str) -> str:
    """
    Convierte:
        NOMBRE_PROYECTO -> Nombre Proyecto
        brecha_1        -> Brecha 1
    """
    return key.replace("_", " ").title()