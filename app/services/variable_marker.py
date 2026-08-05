# app/services/variable_marker.py
"""
Resalta los placeholders [VARIABLE] dentro del .docx para que se distingan
en el editor OnlyOffice, igual que los spans `doc-var` del visor HTML.

El marcado se hace sobre el documento porque una plantilla real tiene cientos
de variables: hacerlo desde el navegador con la API del editor obligaría a una
búsqueda por variable y sería inviable.
"""
import copy
import re
from io import BytesIO
from typing import Dict, Iterator, List, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.parser import VAR_RE
from app.services.sections.heading_parser import _get_heading_level

# Las variables se pintan con sombreado (w:shd), no con el resaltado de Word.
#
# El resaltado (highlight_color) solo admite la paleta fija de Word, y su verde
# es un neón que tapa el texto: con el documento ya redactado no se leía nada.
# El sombreado acepta cualquier color, así que se usan tonos muy lavados que se
# distinguen de un vistazo sin estorbar la lectura.
SHD_VACIA = "FFF6E0"    # ámbar muy lavado: variable sin valor
SHD_RELLENA = "EDF5ED"  # verde muy lavado: variable ya rellenada

# Color de una variable rellenada cuando no se sabe a qué capítulo pertenece
SHD_RELLENA_POR_DEFECTO = SHD_RELLENA

# Cuánto del color del capítulo se conserva al usarlo como fondo. El resto es
# blanco. Con 0.12 el tono se reconoce y el texto se lee sin esfuerzo; subirlo
# devuelve el problema del resaltado que tapaba la letra.
PROPORCION_DE_COLOR = 0.12

# Word limita el nombre de un marcador a 40 caracteres sin espacios
BOOKMARK_MAX_LEN = 40
BOOKMARK_PREFIX = "V_"
# Marcador de título: se numera por su posición de párrafo, que es lo que
# identifica al heading en el índice que consume el frontend
HEADING_PREFIX = "H_"


def _iter_cell_paragraphs(cell) -> Iterator:
    for para in cell.paragraphs:
        yield para
    for table in cell.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table) -> Iterator:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_all_paragraphs(doc: Document) -> Iterator:
    """Cuerpo, tablas (incluidas anidadas), encabezados y pies."""
    for para in doc.paragraphs:
        yield para
    yield from _iter_paragraphs_fuera_del_cuerpo(doc)


def _iter_paragraphs_fuera_del_cuerpo(doc: Document) -> Iterator:
    """Tablas, encabezados y pies: todo lo que no cuenta para el índice de H1."""
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)
    for section in doc.sections:
        for container in (section.header, section.footer):
            for para in container.paragraphs:
                yield para
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _split_segments(text: str) -> List[Tuple[str, str]]:
    """Parte el texto en tramos (contenido, key_o_vacío) según los [PLACEHOLDER]."""
    segments: List[Tuple[str, str]] = []
    pos = 0
    for match in VAR_RE.finditer(text):
        if match.start() > pos:
            segments.append((text[pos:match.start()], ""))
        segments.append((match.group(0), match.group(1)))
        pos = match.end()
    if pos < len(text):
        segments.append((text[pos:], ""))
    return segments


def _set_run_shading(run, fill_hex: str) -> None:
    """
    Pinta el fondo de un run con un color cualquiera.

    python-docx no expone el sombreado de run, así que se monta el elemento
    `w:shd` a mano. Se quita antes cualquier sombreado que traiga el estilo
    original del párrafo para que no se acumulen.
    """
    rpr = run._element.get_or_add_rPr()
    for previo in rpr.findall(qn("w:shd")):
        rpr.remove(previo)

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    rpr.append(shd)


def _quitar_shading(run) -> None:
    """Deja el run sin ningún fondo, venga de donde venga."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return
    for previo in rpr.findall(qn("w:shd")):
        rpr.remove(previo)


def _normalizar_hex(color: str) -> str:
    """Deja el color como lo quiere Word: seis dígitos hex, sin almohadilla."""
    limpio = color.strip().lstrip("#").upper()
    if len(limpio) != 6 or any(c not in "0123456789ABCDEF" for c in limpio):
        raise ValueError(f"Color de variable no válido: {color!r}; se esperaba #RRGGBB")
    return limpio


def _lavar(color_hex: str) -> str:
    """
    Mezcla el color con blanco hasta dejarlo casi apagado.

    Los colores de capítulo están pensados para verse sobre fondo blanco —una
    barrita en el índice, un punto— y a plena intensidad son opacos. Puestos
    detrás del texto lo tapan igual que hacía el verde neón que se quitó.

    Se conserva solo una fracción del color, lo justo para distinguir un
    capítulo de otro sin estorbar la lectura.
    """
    fraccion = PROPORCION_DE_COLOR
    canales = [int(color_hex[i:i + 2], 16) for i in (0, 2, 4)]
    lavados = [round(c * fraccion + 255 * (1 - fraccion)) for c in canales]
    return "".join(f"{c:02X}" for c in lavados)


def _bookmark_name(key: str, used: set) -> str:
    """Nombre de marcador válido para Word, único dentro del documento."""
    safe = re.sub(r"[^A-Za-z0-9_]", "_", key)[: BOOKMARK_MAX_LEN - len(BOOKMARK_PREFIX)]
    name = f"{BOOKMARK_PREFIX}{safe}"

    if name in used:
        suffix = 2
        while f"{name[: BOOKMARK_MAX_LEN - 2]}_{suffix}" in used:
            suffix += 1
        name = f"{name[: BOOKMARK_MAX_LEN - 2]}_{suffix}"

    used.add(name)
    return name


def _wrap_in_bookmark(run, name: str, bookmark_id: int) -> None:
    """Envuelve el run en un marcador para poder saltar hasta él desde el editor."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    run._element.addprevious(start)
    run._element.addnext(end)


def _mark_paragraph(para, values: dict, ctx: dict) -> int:
    """
    Reescribe los runs del párrafo separando cada placeholder en su propio run
    resaltado, y pone un marcador en la primera aparición de cada variable.
    Devuelve cuántas variables marcó.

    Los runs se aplanan al formato del primero — mismo criterio que usa
    `filler._replace_vars_in_paragraph` al rellenar el documento.
    """
    if not para.runs:
        return 0

    full_text = "".join(run.text for run in para.runs)
    if "[" not in full_text:
        return 0

    segments = _split_segments(full_text)
    if not any(key for _, key in segments):
        return 0

    base = para.runs[0]
    base_rpr = base._element.find(qn("w:rPr"))
    base_rpr_copy = copy.deepcopy(base_rpr) if base_rpr is not None else None

    # Vaciar el primer run y eliminar el resto: los segmentos se re-crean debajo
    base.text = ""
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)

    marked = 0
    for text, key in segments:
        value = (values.get(key) or "").strip() if key else ""
        # Con valor se muestra el valor; sin valor, el propio [PLACEHOLDER]
        run = para.add_run(value if (key and value) else text)
        if base_rpr_copy is not None:
            run._element.insert(0, copy.deepcopy(base_rpr_copy))
        if key:
            # El resaltado de Word se limpia por si el documento ya venía
            # marcado con la versión anterior, que usaba el verde neón
            run.font.highlight_color = None

            if not ctx["sombrear"]:
                # Fase de revisión: el documento se entrega limpio. Se quita
                # también cualquier sombreado heredado del estilo, o quedarían
                # restos de la fase anterior.
                _quitar_shading(run)
            elif value:
                # El color sale del capítulo donde vive la variable; si ese
                # capítulo no tiene color asignado se usa el verde lavado
                del_capitulo = ctx["colores_por_h1"].get(ctx["h1_actual"])
                _set_run_shading(run, del_capitulo if del_capitulo else SHD_RELLENA_POR_DEFECTO)
            else:
                _set_run_shading(run, SHD_VACIA)
            marked += 1

            # Solo la primera aparición lleva marcador: es a donde salta el editor
            if key not in ctx["bookmarks"]:
                name = _bookmark_name(key, ctx["used"])
                ctx["bookmarks"][key] = name
                ctx["next_id"] += 1
                _wrap_in_bookmark(run, name, ctx["next_id"])

    return marked


def heading_bookmark(para_index: int) -> str:
    """Nombre del marcador de un título, derivado de su posición de párrafo."""
    return f"{HEADING_PREFIX}{para_index}"


def _mark_headings(doc: Document, ctx: dict) -> int:
    """
    Pone un marcador en cada título del cuerpo. Sin esto el editor no puede
    saltar a un título: buscarlo por texto cae en la tabla de contenidos, que
    repite los mismos textos al principio del documento.
    """
    marked = 0
    for index, para in enumerate(doc.paragraphs):
        if _get_heading_level(para) is None:
            continue
        if not para.text.strip() or not para.runs:
            continue

        name = heading_bookmark(index)
        ctx["next_id"] += 1
        _wrap_in_bookmark(para.runs[0], name, ctx["next_id"])
        ctx["headings"][index] = name
        marked += 1

    return marked


def mark_variables(
    docx_buffer: bytes,
    values: dict | None = None,
    chapter_colors: Dict[int, str] | None = None,
    shading: bool = True,
) -> Tuple[bytes, int, Dict[str, str], Dict[int, str]]:
    """
    Devuelve (docx marcado, variables resaltadas, marcadores de variables,
    marcadores de títulos).

    `values` opcional: las keys con valor se sombrean como rellenadas y el
    resto como vacías.

    `chapter_colors` opcional: índice de párrafo del H1 → color "#RRGGBB" del
    capítulo. Sirve para que en el documento se vea de un vistazo qué capítulo
    rellenó cada cosa, con el mismo color que lleva en el índice del editor.
    Se indexa por H1 y no por variable porque es aquí donde ya se sabe bajo qué
    título vive cada una: el que llama solo tiene que enviar sus capítulos.

    `shading=False` entrega el documento limpio, sin ningún fondo. Es lo que
    hace falta en la fase de revisión: los colores sirven mientras se rellena,
    para ver qué falta y quién puso qué, pero al revisar el contenido estorban
    — hay que leer el informe como lo va a leer quien lo reciba. Los marcadores
    se mantienen igual, así que el índice sigue pudiendo saltar a cada sitio.

    Los marcadores permiten que el editor salte a una variable o a un título.
    """
    doc = Document(BytesIO(docx_buffer))
    values = values or {}

    # Se lavan aquí, una sola vez, en lugar de en cada variable
    colores_por_h1 = {
        int(h1): _lavar(_normalizar_hex(color))
        for h1, color in (chapter_colors or {}).items()
    }

    # Los ids arrancan altos para no chocar con los marcadores propios del docx
    ctx = {
        "bookmarks": {},
        "headings": {},
        "used": set(),
        "next_id": 90000,
        "colores_por_h1": colores_por_h1,
        "h1_actual": None,
        "sombrear": shading,
    }

    # Primero los títulos: marcar variables reescribe runs y desplaza posiciones
    _mark_headings(doc, ctx)

    total = 0

    # Cuerpo: se lleva la cuenta del H1 vigente, que es el mismo criterio con el
    # que el parser asigna `h1_index` a cada variable
    for index, para in enumerate(doc.paragraphs):
        if _get_heading_level(para) == 1:
            ctx["h1_actual"] = index
        total += _mark_paragraph(para, values, ctx)

    # Tablas, encabezados y pies heredan el último capítulo del cuerpo
    for para in _iter_paragraphs_fuera_del_cuerpo(doc):
        total += _mark_paragraph(para, values, ctx)

    out = BytesIO()
    doc.save(out)
    return out.getvalue(), total, ctx["bookmarks"], ctx["headings"]


def list_bookmarks(docx_buffer: bytes) -> Dict[str, str]:
    """Mapa key → nombre de marcador de variable, sin devolver el documento."""
    _, _, bookmarks, _ = mark_variables(docx_buffer)
    return bookmarks
